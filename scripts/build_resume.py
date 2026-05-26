"""Generate Jordan Henning's resume in active variants:

  Universal:
    Jordan-Henning-Resume.docx           Full deep-dive (~4 pages)
    Jordan-Henning-Resume-Brief.docx     Recruiter-scan brief (~2 pages)

  Role-tailored briefs (~2 pages each, custom Executive Summary for the role):
    Jordan-Henning-Resume-Service-Ops.docx   Federal Service Delivery / Operations
    Jordan-Henning-Resume-Federal-IT.docx    Federal IT Leadership
    Jordan-Henning-Resume-Program-PM.docx    Project / Program Management

Run:
  python scripts/build_resume.py
  python scripts/build_resume.py --pdf      # also produce PDFs (Word required)
"""

from __future__ import annotations

import argparse
import subprocess
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

# ---------------------------------------------------------------------------
# Theme
# ---------------------------------------------------------------------------

FONT = "Calibri"
NAVY = RGBColor(0x0B, 0x22, 0x4D)
ACCENT = RGBColor(0x1F, 0x4E, 0x8C)
TEXT = RGBColor(0x1A, 0x1F, 0x2C)
MUTED = RGBColor(0x55, 0x60, 0x77)
RULE = RGBColor(0xC9, 0xD3, 0xE2)


def set_run(run, *, size=11, bold=False, italic=False, color=TEXT, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)


def add_para(doc, *, space_before=0, space_after=2, alignment=None, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if alignment is not None:
        p.alignment = alignment
    if indent is not None:
        pf.left_indent = Inches(indent)
    return p


def add_run(p, text, **kwargs):
    run = p.add_run(text)
    set_run(run, **kwargs)
    return run


def horizontal_rule(p, color=RULE):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    hex_color = "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2])
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_header(doc, title, *, space_before=10, space_after=4):
    p = add_para(doc, space_before=space_before, space_after=space_after)
    add_run(p, title.upper(), size=11, bold=True, color=NAVY)
    rule = add_para(doc, space_before=0, space_after=6)
    horizontal_rule(rule)
    return p


def add_bullet(doc, runs, indent=0.25, space_after=3):
    p = add_para(doc, space_after=space_after, indent=indent)
    pf = p.paragraph_format
    pf.first_line_indent = Inches(-0.18)
    pf.left_indent = Inches(indent)
    add_run(p, "•  ", size=11, color=ACCENT, bold=True)
    if isinstance(runs, str):
        add_run(p, runs, size=11)
    else:
        for text, kwargs in runs:
            add_run(p, text, **kwargs)
    return p


def add_role_block(
    doc,
    *,
    role,
    org,
    dates,
    location=None,
    lead=None,
    bullets=None,
    space_before=8,
):
    role_p = add_para(doc, space_before=space_before, space_after=0)
    add_run(role_p, role, size=12, bold=True, color=TEXT)
    add_run(role_p, "    ·  ", size=11, color=MUTED)
    add_run(role_p, org, size=12, color=ACCENT, bold=True)

    meta_p = add_para(doc, space_before=0, space_after=4)
    add_run(meta_p, dates, size=10, italic=True, color=MUTED)
    if location:
        add_run(meta_p, "    ·  ", size=10, color=MUTED)
        add_run(meta_p, location, size=10, italic=True, color=MUTED)

    if lead:
        lead_p = add_para(doc, space_after=4)
        add_run(lead_p, lead, size=11, italic=True, color=TEXT)

    if bullets:
        for b in bullets:
            add_bullet(doc, b)


# ---------------------------------------------------------------------------
# Shared content
# ---------------------------------------------------------------------------

CONTACT_PARTS = [
    "York, PA",
    "JordanHenning32@yahoo.com",
    "330-280-0642",
    "jordanhenning.com",
    "linkedin.com/in/jordan-henning",
    "github.com/jordanhenning32",
]

HEADER_TITLE = "Federal Services Delivery Executive  ·  17+ years SSA  ·  FAC-P/PM"


# ---------------------------------------------------------------------------
# Executive Summaries — one per variant
# ---------------------------------------------------------------------------

def summary_full():
    return [
        ("Federal services delivery executive who codes the AI. ", dict(size=11)),
        ("Seventeen years of federal IT at the Social Security Administration", dict(size=11, bold=True)),
        (", culminating in three years as Branch Chief — manager-of-managers over ", dict(size=11)),
        ("12 direct reports + 340 field IT staff across 170 nationwide Hearings Offices, serving 7,000+ SSA employees", dict(size=11, bold=True)),
        (" at ", dict(size=11)),
        ("99.9% system availability", dict(size=11, bold=True)),
        (", with multi-year ", dict(size=11)),
        ("$200M+ Agile IT portfolios", dict(size=11, bold=True)),
        (" delivered on-time and under-budget. ", dict(size=11)),
        ("Public Trust High Risk clearance", dict(size=11, bold=True)),
        (" held throughout SSA tenure (2008–2025), eligible to reinstate. ", dict(size=11)),
        ("FAC-P/PM certified", dict(size=11, bold=True)),
        ("; familiar with GSA MAS Schedule, STARS III, 8(a)/SDVOSB/HUBZone, IDIQ/BPA, FedRAMP, ATO governance, and PA Invitation-to-Qualify procurement. Now ", dict(size=11)),
        ("Chief Growth Officer at Quadratic Digital", dict(size=11, bold=True)),
        (" (14-person federal services firm), where I personally designed, coded, and operate two production multi-agent AI systems: ", dict(size=11)),
        ("RFP Factory", dict(size=11, bold=True)),
        (" (multi-agent proposal automation, 40-hour federal proposal cycles compressed to 2 hours — 20× faster) and ", dict(size=11)),
        ("Futures Bot", dict(size=11, bold=True)),
        (" (autonomous multi-agent trading, live since Feb 2026, 62% win rate over 500+ trades). Combat-tested under three tours with the 101st Airborne (Bronze Star, Purple Heart). Targeting ", dict(size=11)),
        ("VP/Director Federal Services Delivery, GM Federal AI Services, or Senior Director Federal Programs", dict(size=11, bold=True)),
        (" — where federal-scale uptime, $200M+ portfolio discipline, procurement fluency, and hands-on AI capability all matter in the same role.", dict(size=11)),
    ]


def summary_brief():
    return [
        ("Federal services delivery executive who codes the AI. ", dict(size=11)),
        ("17 years federal IT at SSA", dict(size=11, bold=True)),
        (" — most recently three years as ", dict(size=11)),
        ("Branch Chief: manager-of-managers over 12 direct + 340 field IT across 170 offices serving 7,000+ employees (99.9% availability, $200M+ portfolio)", dict(size=11, bold=True)),
        (". FAC-P/PM certified. Public Trust cleared. Layered with two production multi-agent AI systems I designed, coded, and operate myself: ", dict(size=11)),
        ("RFP Factory", dict(size=11, bold=True)),
        (" (multi-agent proposal automation, 20× cycle compression at Quadratic Digital) and ", dict(size=11)),
        ("Futures Bot", dict(size=11, bold=True)),
        (" (autonomous multi-agent trading, live since Feb 2026, 62% win rate over ~500 trades). Combat veteran, 101st Airborne (Bronze Star, Purple Heart). Targeting ", dict(size=11)),
        ("VP/Director Federal Services Delivery", dict(size=11, bold=True)),
        (" or ", dict(size=11)),
        ("GM Federal AI Services", dict(size=11, bold=True)),
        (".", dict(size=11)),
    ]


def summary_service_ops():
    return [
        ("Federal services delivery executive", dict(size=11, bold=True)),
        (" with 17 years at SSA, culminating in three years as ", dict(size=11)),
        ("Branch Chief sustaining 99.9% availability across 170 nationwide offices", dict(size=11, bold=True)),
        (" for mission-critical platforms serving millions of users. Reduced system downtime 20% through ITIL-driven monitoring, dashboards, and continual service improvement. Manager-of-managers over 12 direct reports + 340 field IT staff serving 7,000+ Hearings Office employees, with multimillion-dollar budgeting, full vendor lifecycle management, ", dict(size=11)),
        ("FedRAMP / ATO governance", dict(size=11, bold=True)),
        (", and CIO-level cross-functional partnership. That service-delivery discipline is now layered with hands-on multi-agent AI fluency — the foundation for AI-augmented federal service operations. Public Trust cleared. Combat veteran, 101st Airborne (Bronze Star, Purple Heart). Targeting ", dict(size=11)),
        ("VP/Director Federal Services Delivery, GM Federal AI Services, or Senior Director Federal Programs", dict(size=11, bold=True)),
        (".", dict(size=11)),
    ]


def summary_federal_it():
    return [
        ("Federal IT executive", dict(size=11, bold=True)),
        (" with 17 years at the Social Security Administration, culminating in three years as Branch Chief — manager-of-managers over 12 direct reports and 340 field IT staff across 170 nationwide Hearings Offices serving 7,000+ employees. ", dict(size=11)),
        ("Public Trust High Risk clearance", dict(size=11, bold=True)),
        (" held throughout SSA tenure (2008–2025), eligible to reinstate. ", dict(size=11)),
        ("FAC-P/PM certified", dict(size=11, bold=True)),
        ("; familiar with GSA MAS Schedule, STARS III, 8(a)/SDVOSB/HUBZone, IDIQ/BPA, FedRAMP, and ATO governance. SSA Commissioner Award (2021) for emergency COVID-19 service delivery. 101st Airborne combat veteran (Bronze Star, Purple Heart). Now layered with hands-on multi-agent AI fluency for federal AI service delivery. Targeting ", dict(size=11)),
        ("senior federal IT leadership roles", dict(size=11, bold=True)),
        (" — federal contractors, federal services firms, defense AI services, or regulated commercial roles where federal experience and clearance are core.", dict(size=11)),
    ]


def summary_program_pm():
    return [
        ("FAC-P/PM-certified federal program leader", dict(size=11, bold=True)),
        (". Nine years of federal IT leadership at SSA — five years as IT Project Manager directing ", dict(size=11)),
        ("$200M+ Agile portfolios on-time and under-budget", dict(size=11, bold=True)),
        (", three years as Branch Chief manager-of-managers over 12 direct + 340 field IT staff across 170 offices serving 7,000+ Hearings employees. Named transformation work: agency-wide BI platform implementation (Tableau + WebFocus), 7-ODS consolidation into a single Appeals Database, centralized print services consolidation (multimillion-dollar savings), ", dict(size=11)),
        ("VA cross-agency disability-evidence partnership", dict(size=11, bold=True)),
        (", and the emergency COVID-19 medical document upload system on MySSA (SSA Commissioner Award). 100% direct-report retention; 4 of 4 mentees promoted into PM leadership. Now layered with hands-on multi-agent AI fluency for AI-augmented program offices. Combat veteran, 101st Airborne (Bronze Star, Purple Heart). Targeting ", dict(size=11)),
        ("senior PM/PgM roles at federal contractors and federal services firms", dict(size=11, bold=True)),
        (" where federal-scale discipline meets modern delivery.", dict(size=11)),
    ]


# ---------------------------------------------------------------------------
# Outcomes
# ---------------------------------------------------------------------------

OUTCOMES_FULL = [
    [
        ("Manager-of-managers as Branch Chief", dict(size=11, bold=True)),
        (" — 12 direct reports + 340 field IT across 170 nationwide offices serving 7,000+ employees at ", dict(size=11)),
        ("99.9% availability", dict(size=11, bold=True)),
    ],
    [
        ("$200M+ Agile IT portfolio", dict(size=11, bold=True)),
        (" delivered on-time, under-budget over multi-year federal programs", dict(size=11)),
    ],
    [
        ("FedRAMP-certified delivery posture", dict(size=11, bold=True)),
        (" · ATO governance across SSA mission-critical systems · GSA MAS, STARS III, 8(a)/SDVOSB/HUBZone, IDIQ/BPA pursuit experience", dict(size=11)),
    ],
    [
        ("20× cycle-time compression", dict(size=11, bold=True)),
        (" on enterprise proposal workflow via personally-built multi-agent AI system (40 labor hours → 2)", dict(size=11)),
    ],
    [
        ("100% direct-report retention", dict(size=11, bold=True)),
        (" over 3 years as Branch Chief; ", dict(size=11)),
        ("4 of 4 mentees", dict(size=11, bold=True)),
        (" promoted into PM leadership", dict(size=11)),
    ],
    [
        ("20% system downtime reduction", dict(size=11, bold=True)),
        (" through ITIL-driven monitoring, dashboards, and continual service improvement", dict(size=11)),
    ],
    [
        ("Multimillion-dollar agency-wide cost savings", dict(size=11, bold=True)),
        (" via centralized print consolidation and vendor renegotiation", dict(size=11)),
    ],
    [
        ("SSA Commissioner Award", dict(size=11, bold=True)),
        (" for emergency COVID-19 service delivery during nationwide office closures", dict(size=11)),
    ],
]

OUTCOMES_BRIEF = OUTCOMES_FULL[:5]


# ---------------------------------------------------------------------------
# Production AI Builds
# ---------------------------------------------------------------------------

def add_ai_builds(doc, *, brief=False):
    add_section_header(doc, "Production AI Builds (Hands-on)")

    title_p = add_para(doc, space_before=2, space_after=2)
    add_run(title_p, "RFP Factory", size=12, bold=True, color=TEXT)
    add_run(title_p, "  —  ", size=11, color=MUTED)
    add_run(title_p, "Multi-agent proposal automation", size=11, italic=True, color=ACCENT)
    add_run(title_p, "  ·  In production at Quadratic Digital", size=10, color=MUTED)

    body_p = add_para(doc, space_after=4)
    if brief:
        runs = [
            ("Multi-agent system that ingests RFP packages and produces SME-ready proposal drafts across an Anthropic + OpenAI + Gemini stack. ", dict(size=11)),
            ("Compresses a 40-hour proposal task into 2 labor hours of human review — a 20× reduction.", dict(size=11, bold=True)),
            (" Designed, coded, operated.", dict(size=11)),
        ]
    else:
        runs = [
            ("End-to-end system that ingests federal and commercial RFP packages and produces polished, SME-ready proposal drafts. Specialized agents handle research, strategy, drafting, and compliance review in parallel across a multi-provider LLM stack (Anthropic + OpenAI + Gemini). ", dict(size=11)),
            ("Compresses a 40-hour proposal task into 2 labor hours of human review — a 20× reduction.", dict(size=11, bold=True)),
            (" Designed, coded, and operated by Jordan. Stack: Python 3.14, SQLAlchemy 2.0, Alembic, NiceGUI, ThreadPoolExecutor, prompt engineering at the agent-system level.", dict(size=11)),
        ]
    for text, kwargs in runs:
        add_run(body_p, text, **kwargs)

    title_p = add_para(doc, space_before=4, space_after=2)
    add_run(title_p, "Futures Bot", size=12, bold=True, color=TEXT)
    add_run(title_p, "  —  ", size=11, color=MUTED)
    add_run(title_p, "Autonomous multi-agent futures trading", size=11, italic=True, color=ACCENT)
    add_run(title_p, "  ·  Live since Feb 2026  ·  62% win rate over ~500 trades", size=10, color=MUTED)

    body_p = add_para(doc, space_after=4)
    if brief:
        runs = [
            ("Personal R&D platform: ", dict(size=11)),
            ("five-agent committee (Technical Analyst, Bull, Bear, Risk Manager, Devil's Advocate)", dict(size=11, bold=True)),
            (" running across OpenAI, Claude, and Gemini. Trades futures markets autonomously using strategies built on twelve years of personal markets research. ~8 trades per day across the 24-hour cycle.", dict(size=11)),
        ]
    else:
        runs = [
            ("Personal R&D platform: a multi-agent system that monitors futures markets, generates signals, manages risk, and executes trades autonomously, using proprietary strategies developed from twelve years of personal markets research. Live since February 2026 — ~8 trades per day across the 24-hour cycle, 62% win rate over the first ~500 trades, net-positive PnL. ", dict(size=11)),
            ("Five-agent committee (Technical Analyst, Bull, Bear, Risk Manager, Devil's Advocate)", dict(size=11, bold=True)),
            (" running across OpenAI, Claude, and Gemini. Built end-to-end as a research platform for pressure-testing multi-agent design patterns under conditions where mistakes cost real money.", dict(size=11)),
        ]
    for text, kwargs in runs:
        add_run(body_p, text, **kwargs)


# ---------------------------------------------------------------------------
# Role bullets (shared)
# ---------------------------------------------------------------------------

CGO_BULLETS_FULL = [
    [
        ("Designed and shipped ", dict(size=11)),
        ("RFP Factory", dict(size=11, bold=True)),
        (", a multi-agent proposal automation system that compresses 40-hour proposal cycles into 2 hours of human review — a 20× reduction no traditional GTM team can match.", dict(size=11)),
    ],
    [
        ("Subcontracted multiple roles to ", dict(size=11)),
        ("Nava on CMS modernization programs", dict(size=11, bold=True)),
        (" — embedding Quadratic into a marquee civic-tech delivery prime supporting healthcare-focused federal modernization work.", dict(size=11)),
    ],
    [
        ("Won ", dict(size=11)),
        ("two PA Invitation-to-Qualify (ITQ) vehicles", dict(size=11, bold=True)),
        (" and qualified Quadratic into the ", dict(size=11)),
        ("PA Small Disadvantaged Business program", dict(size=11, bold=True)),
        ("; pursuing capture across GSA MAS Schedule, STARS III, 8(a)/SDVOSB/HUBZone vehicles, and IDIQ/BPA structures.", dict(size=11)),
    ],
    [
        ("Operates a ", dict(size=11)),
        ("FedRAMP-certified delivery posture", dict(size=11, bold=True)),
        (" — AI-services builds run alongside ATO-aligned controls so the federal compliance baseline is the starting point, not an afterthought.", dict(size=11)),
    ],
    "Executing cross-sector growth strategy across federal and commercial markets — pricing architecture, solution scoping, and offer-portfolio development.",
    "Winning high-value strategic contracts and building executive partnerships in highly regulated environments where service delivery and AI maturity have to coexist.",
]

CGO_BULLETS_BRIEF = [
    [
        ("Designed and shipped ", dict(size=11)),
        ("RFP Factory", dict(size=11, bold=True)),
        (" — multi-agent proposal automation that compresses 40-hour cycles into 2 hours (20× reduction).", dict(size=11)),
    ],
    [
        ("Subcontracted Quadratic into ", dict(size=11)),
        ("Nava on CMS modernization", dict(size=11, bold=True)),
        ("; won ", dict(size=11)),
        ("two PA ITQ vehicles", dict(size=11, bold=True)),
        ("; qualified into PA Small Disadvantaged Business program; pursuing capture across GSA MAS, STARS III, 8(a)/SDVOSB/HUBZone, IDIQ/BPA.", dict(size=11)),
    ],
    [
        ("FedRAMP-certified delivery posture", dict(size=11, bold=True)),
        (" with ATO-aligned controls — federal compliance baseline as the starting point, not an afterthought.", dict(size=11)),
    ],
    "Executing cross-sector growth strategy across federal and commercial markets; winning strategic contracts in regulated environments where service delivery and AI maturity have to coexist.",
]

BRANCH_BULLETS_FULL = [
    "Owned end-to-end service strategy, incident/problem management, escalation handling, root-cause analysis, and SLA/KPI governance for a 24/7 federal IT operation.",
    [
        ("Reduced system downtime 20%", dict(size=11, bold=True)),
        (" through enhanced monitoring frameworks, dashboards, proactive prevention, and ITIL-driven continual service improvement.", dict(size=11)),
    ],
    "Directed multimillion-dollar budgeting, procurement, vendor relationships, and full lifecycle infrastructure refresh — delivered on-time, under-budget.",
    [
        ("Championed cross-functional collaboration with CIO-level leadership, engineering, security, and stakeholder teams; ensured regulatory compliance and audit readiness — including ", dict(size=11)),
        ("FedRAMP and ATO governance", dict(size=11, bold=True)),
        (" across mission-critical systems.", dict(size=11)),
    ],
    [
        ("Oversaw staffing, resource planning, and performance optimization to deliver consistent customer outcomes; ", dict(size=11)),
        ("100% direct-report retention", dict(size=11, bold=True)),
        (" across three-year tenure.", dict(size=11)),
    ],
]

BRANCH_BULLETS_BRIEF = [
    [
        ("Reduced system downtime 20%", dict(size=11, bold=True)),
        (" through enhanced monitoring, dashboards, and ITIL-driven continual service improvement.", dict(size=11)),
    ],
    [
        ("Owned end-to-end service strategy, incident/escalation management, and SLA/KPI governance for a 24/7 federal operation; ", dict(size=11)),
        ("FedRAMP / ATO governance", dict(size=11, bold=True)),
        (" across mission-critical systems; multimillion-dollar budgeting, procurement, and lifecycle refresh.", dict(size=11)),
    ],
    [
        ("100% direct-report retention", dict(size=11, bold=True)),
        (" across three-year tenure as manager-of-managers (12 direct, 340 indirect reports).", dict(size=11)),
    ],
]

ITPM_BULLETS_FULL = [
    [
        ("Led ", dict(size=11)),
        ("centralized print services consolidation", dict(size=11, bold=True)),
        (" — single-vendor architecture and renegotiated contracts saving the agency millions of dollars annually.", dict(size=11)),
    ],
    [
        ("Implemented ", dict(size=11)),
        ("agency-wide BI platform", dict(size=11, bold=True)),
        (" — Tableau + WebFocus selection, integration, training, and ongoing maintenance, enabling data-driven decisions across the SSA.", dict(size=11)),
    ],
    [
        ("Consolidated ", dict(size=11)),
        ("seven legacy Operational Data Stores", dict(size=11, bold=True)),
        (" into one modern Appeals Database — improved accuracy, speed, and cost efficiency at federal scale.", dict(size=11)),
    ],
    [
        ("Partnered with the ", dict(size=11)),
        ("U.S. Department of Veterans Affairs", dict(size=11, bold=True)),
        (" in 2016 to enable secure cross-agency disability-evidence sharing — surfacing VA medical records earlier in SSA's disability-claim process to accelerate veteran benefit approvals.", dict(size=11)),
    ],
    [
        ("Spearheaded the ", dict(size=11)),
        ("emergency COVID-19 medical document upload system", dict(size=11, bold=True)),
        (" on MySSA, keeping disability-claim processing uninterrupted during nationwide office closures. Earned the SSA Commissioner Award.", dict(size=11)),
    ],
    "Leveraged ServiceNow for ITSM, workflow automation, incident tracking, and self-service at enterprise scale.",
    [
        ("Mentored 4 project managers", dict(size=11, bold=True)),
        (", all of whom went on to lead their own programs successfully.", dict(size=11)),
    ],
]

ITPM_BULLETS_BRIEF = [
    [
        ("Spearheaded the ", dict(size=11)),
        ("emergency COVID-19 medical document upload system", dict(size=11, bold=True)),
        (" on MySSA — earned the SSA Commissioner Award. Partnered with the VA in 2016 on cross-agency disability-evidence sharing to accelerate veteran benefit approvals.", dict(size=11)),
    ],
    [
        ("Led ", dict(size=11)),
        ("centralized print services consolidation", dict(size=11, bold=True)),
        (" — single-vendor architecture saving the agency millions annually; implemented agency-wide BI platform (Tableau + WebFocus); consolidated 7 legacy ODS into one Appeals Database.", dict(size=11)),
    ],
    [
        ("Mentored 4 project managers", dict(size=11, bold=True)),
        (", all promoted into program leadership.", dict(size=11)),
    ],
]

ANALYST_BULLETS_FULL = [
    "Reduced system downtime 20% through proactive analysis, troubleshooting, process redesign, and enhanced monitoring across high-volume beneficiary systems.",
    "Delivered Tier-3 IT support and systems administration across multiple field offices; increased network uptime and staff productivity through targeted training and issue resolution.",
    "Contributed to early MySSA iterations, leveraging field-office Claims Representative experience to inform the SSA's transition of services online.",
]

ANALYST_BULLETS_BRIEF = [
    "Reduced downtime 20% across high-volume beneficiary systems; delivered Tier-3 IT support across multiple field offices; contributed to early MySSA iterations.",
]

ARMY_BULLETS_FULL = [
    [
        ("Three combat tours (Iraq, Afghanistan); awarded the ", dict(size=11)),
        ("Bronze Star", dict(size=11, bold=True)),
        (" and ", dict(size=11)),
        ("Purple Heart", dict(size=11, bold=True)),
        (".", dict(size=11)),
    ],
    "Developed executive-level leadership, adaptability, resilience, and performance under extreme pressure — directly transferable to high-stakes federal service delivery and incident-command escalations.",
]

RECOGNITION = [
    [
        ("SSA Commissioner Award", dict(size=11, bold=True)),
        ("  ·  2021  ·  Presented by Acting Commissioner Kilolo Kijakazi for COVID-19 emergency document upload delivery", dict(size=11)),
    ],
    [
        ("Bronze Star", dict(size=11, bold=True)),
        ("  ·  2006  ·  U.S. Army (Iraq combat tour)", dict(size=11)),
    ],
    [
        ("Purple Heart", dict(size=11, bold=True)),
        ("  ·  2006  ·  U.S. Army", dict(size=11)),
    ],
    [
        ("FAC-P/PM Certification", dict(size=11, bold=True)),
        ("  ·  2020  ·  Federal Acquisition Institute", dict(size=11)),
    ],
    [
        ("Public Trust Clearance · High Risk Tier", dict(size=11, bold=True)),
        ("  ·  Held throughout 17-year SSA tenure (2008–2025)  ·  Eligible to reinstate", dict(size=11)),
    ],
]

EDUCATION = [
    [
        ("Master of Business Administration (MBA)", dict(size=11, bold=True)),
        ("  ·  Malone University  ·  2012", dict(size=11)),
    ],
    [
        ("B.A., Computer Information Systems", dict(size=11, bold=True)),
        ("  ·  Kent State University  ·  2007", dict(size=11)),
    ],
]


# ---------------------------------------------------------------------------
# Build doc
# ---------------------------------------------------------------------------

def build_doc(*, brief=False, summary_runs=None) -> Document:
    """summary_runs: callable that returns the runs list for the Executive Summary.
    If None, defaults to summary_brief() or summary_full() based on `brief`."""
    if summary_runs is None:
        summary_runs = summary_brief if brief else summary_full

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.1 if brief else 1.4)
        section.bottom_margin = Cm(1.1 if brief else 1.4)
        section.left_margin = Cm(1.5 if brief else 1.6)
        section.right_margin = Cm(1.5 if brief else 1.6)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT

    # Header
    name_p = add_para(doc, space_after=2)
    add_run(name_p, "Jordan Henning", size=24 if brief else 26, bold=True, color=NAVY)

    title_p = add_para(doc, space_after=2)
    add_run(
        title_p,
        HEADER_TITLE,
        size=11 if brief else 12,
        color=ACCENT,
        bold=True,
    )

    contact_p = add_para(doc, space_after=4)
    for i, part in enumerate(CONTACT_PARTS):
        if i > 0:
            add_run(contact_p, "  ·  ", size=10, color=MUTED)
        add_run(contact_p, part, size=10, color=MUTED)

    rule_p = add_para(doc, space_before=2, space_after=6 if brief else 8)
    horizontal_rule(rule_p, color=NAVY)

    # Summary
    add_section_header(doc, "Executive Summary")
    summary_p = add_para(doc, space_after=4)
    for text, kwargs in summary_runs():
        add_run(summary_p, text, **kwargs)

    # Outcomes
    add_section_header(doc, "Headline Outcomes")
    outcomes = OUTCOMES_BRIEF if brief else OUTCOMES_FULL
    for runs in outcomes:
        add_bullet(doc, runs, space_after=2 if brief else 3)

    # Production AI Builds (full only — brief mentions inline in summary)
    if not brief:
        add_ai_builds(doc, brief=False)

    # Experience
    add_section_header(doc, "Experience")

    add_role_block(
        doc,
        role="Chief Growth Officer",
        org="Quadratic Digital",
        dates="April 2025 – Present",
        lead=None if brief else (
            "Driving growth at a 14-person services firm serving state, federal, and prime-subcontract clients — "
            "pairing proven go-to-market discipline with the multi-agent AI tooling I build hands-on."
        ),
        bullets=CGO_BULLETS_BRIEF if brief else CGO_BULLETS_FULL,
    )

    add_role_block(
        doc,
        role="Branch Chief, Hearings Office IT Oversight",
        org="Social Security Administration",
        dates="January 2022 – April 2025",
        location="Baltimore, MD",
        lead=None if brief else (
            "Manager-of-managers from a 12-direct HQ team overseeing 340 field IT staff across 170 nationwide "
            "Hearings Offices, delivering mission-critical IT services to 7,000+ Hearings Office employees at 99.9% availability."
        ),
        bullets=BRANCH_BULLETS_BRIEF if brief else BRANCH_BULLETS_FULL,
    )

    add_role_block(
        doc,
        role="IT Project Manager (FAC-P/PM Certified)",
        org="Social Security Administration",
        dates="September 2016 – January 2022",
        location="Baltimore, MD",
        lead=None if brief else (
            "Directed multi-year Agile IT service and infrastructure programs supporting nationwide claims and "
            "hearings systems with portfolios exceeding $200M."
        ),
        bullets=ITPM_BULLETS_BRIEF if brief else ITPM_BULLETS_FULL,
    )

    add_role_block(
        doc,
        role="Systems Analyst · Area Systems Coordinator · Claims Representative",
        org="Social Security Administration",
        dates="August 2008 – September 2016",
        location="Baltimore, MD",
        bullets=ANALYST_BULLETS_BRIEF if brief else ANALYST_BULLETS_FULL,
    )

    if brief:
        # Combine Army + Software Dev into a single "Earlier" line
        earlier_p = add_para(doc, space_before=6, space_after=2)
        add_run(earlier_p, "Earlier:  ", size=11, bold=True, color=NAVY)
        add_run(earlier_p, "U.S. Army, 101st Airborne Division", size=11, bold=True)
        add_run(
            earlier_p,
            " (2001–2009) — three combat tours (Iraq, Afghanistan), Bronze Star, Purple Heart.  ·  ",
            size=11,
        )
        add_run(
            earlier_p,
            "Software Developer at MTD Products (2007–2008).",
            size=11,
        )
    else:
        add_role_block(
            doc,
            role="Infantry Soldier · Transportation Specialist",
            org="U.S. Army, 101st Airborne Division",
            dates="January 2001 – January 2009",
            bullets=ARMY_BULLETS_FULL,
        )
        add_role_block(
            doc,
            role="Software Developer",
            org="MTD Products",
            dates="January 2007 – August 2008",
            location="Valley City, OH",
            bullets=[
                "Developed manufacturing and inventory-management applications in an Agile environment; delivered updates ahead of schedule with emphasis on reliability and operational efficiency.",
            ],
        )

    # Recognition & Credentials
    add_section_header(doc, "Recognition & Credentials")
    for runs in RECOGNITION:
        add_bullet(doc, runs, space_after=2 if brief else 3)

    # Education
    add_section_header(doc, "Education")
    for runs in EDUCATION:
        add_bullet(doc, runs, space_after=2)

    if brief:
        vet_p = add_para(doc, space_before=4, space_after=2)
        add_run(vet_p, "Veteran status:  ", size=11, bold=True, color=NAVY)
        add_run(
            vet_p,
            "Service-connected disabled veteran  ·  Veterans Preference eligible.",
            size=11,
        )
    else:
        add_section_header(doc, "Veteran Status")
        vet_p = add_para(doc, space_after=4)
        add_run(
            vet_p,
            "Service-connected disabled veteran  ·  Veterans Preference eligible",
            size=11,
        )

        # References
        add_section_header(doc, "References")
        ref_intro = add_para(doc, space_after=4)
        add_run(
            ref_intro,
            "Available after first conversation, per professional courtesy:",
            size=11,
            italic=True,
            color=MUTED,
        )
        for r in [
            "1 Director-level reference",
            "3 SSA peer Branch Chief references",
            "Multiple direct reports across SSA + Quadratic Digital",
        ]:
            add_bullet(doc, r)

        # What I'm Looking For
        add_section_header(doc, "What I'm Looking For")
        best_p = add_para(doc, space_after=3)
        add_run(best_p, "Best fit:  ", size=11, bold=True, color=NAVY)
        add_run(
            best_p,
            (
                "VP/Director Federal Services Delivery, GM Federal AI Services, or Senior Director Federal Programs "
                "at federal primes and services firms — where 99.9% uptime, $200M+ portfolio scale, and hands-on AI "
                "capability all matter in the same role. Also: GM, VP, or Practice Director at emerging federal AI "
                "services firms where someone who can both win federal contracts and sponsor multi-agent AI delivery "
                "is the job, not a side capability."
            ),
            size=11,
        )

        less_p = add_para(doc, space_after=4)
        add_run(less_p, "Less of a fit:  ", size=11, bold=True, color=NAVY)
        add_run(
            less_p,
            (
                "VP Product or VP Engineering at commercial tech (different level/tenure pattern)  ·  "
                "IC engineering roles  ·  pre-revenue seed-stage where federal credibility doesn't carry weight  ·  "
                "roles where AI is a marketing layer rather than a real operational change."
            ),
            size=11,
        )

    # Brief footer
    if brief:
        footer_p = add_para(doc, space_before=8, space_after=2)
        add_run(
            footer_p,
            "Full version + writing samples + chatbot Q&A:  ",
            size=10,
            italic=True,
            color=MUTED,
        )
        add_run(footer_p, "jordanhenning.com", size=10, italic=True, color=ACCENT, bold=True)

    return doc


# ---------------------------------------------------------------------------
# Variants
# ---------------------------------------------------------------------------

VARIANTS = [
    {
        "filename": "Jordan-Henning-Resume.docx",
        "brief": False,
        "summary_runs": summary_full,
        "label": "General — Full deep-dive (~4 pages)",
    },
    {
        "filename": "Jordan-Henning-Resume-Brief.docx",
        "brief": True,
        "summary_runs": summary_brief,
        "label": "General — 2-page brief",
    },
    {
        "filename": "Jordan-Henning-Resume-Service-Ops.docx",
        "brief": True,
        "summary_runs": summary_service_ops,
        "label": "Federal Service Delivery / Operations",
    },
    {
        "filename": "Jordan-Henning-Resume-Federal-IT.docx",
        "brief": True,
        "summary_runs": summary_federal_it,
        "label": "Federal IT Leadership",
    },
    {
        "filename": "Jordan-Henning-Resume-Program-PM.docx",
        "brief": True,
        "summary_runs": summary_program_pm,
        "label": "Project / Program Management",
    },
]


# Files left over from deprecated variants (cleaned up at build time so a stale
# PDF can't be downloaded after the picker stops listing it).
DEPRECATED_FILES = [
    "Jordan-Henning-Resume-VP-Senior.docx",
    "Jordan-Henning-Resume-VP-Senior.pdf",
    "Jordan-Henning-Resume-AI-Leadership.docx",
    "Jordan-Henning-Resume-AI-Leadership.pdf",
    "Jordan-Henning-Resume-Product.docx",
    "Jordan-Henning-Resume-Product.pdf",
    "Jordan-Henning-Resume-Growth.docx",
    "Jordan-Henning-Resume-Growth.pdf",
]


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def docx_to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    ps = (
        "$word = New-Object -ComObject Word.Application;"
        "$word.Visible = $false;"
        f"$doc = $word.Documents.Open('{docx_path}');"
        f"$doc.SaveAs([ref] '{pdf_path}', [ref] 17);"
        "$doc.Close();"
        "$word.Quit();"
        "[System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null;"
    )
    subprocess.run(["powershell", "-NonInteractive", "-NoProfile", "-Command", ps], check=True)
    return pdf_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--pdf", action="store_true", help="Also produce PDF (requires Word)")
    parser.add_argument(
        "--out",
        default=str(Path(__file__).resolve().parent.parent / "public" / "resumes"),
        help="Output directory",
    )
    args = parser.parse_args()

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Remove deprecated files first so a stale PDF can't linger if regen fails.
    for stale in DEPRECATED_FILES:
        stale_path = out_dir / stale
        if stale_path.exists():
            stale_path.unlink()
            print(f"  REMOVED: {stale_path.name}  (deprecated variant)")

    for variant in VARIANTS:
        doc = build_doc(brief=variant["brief"], summary_runs=variant["summary_runs"])
        out_path = out_dir / variant["filename"]
        doc.save(str(out_path))
        size_kb = out_path.stat().st_size // 1024
        print(f"  WROTE: {out_path.name}  ({size_kb} KB)  — {variant['label']}")

        if args.pdf:
            pdf_path = docx_to_pdf(out_path)
            pdf_size = pdf_path.stat().st_size // 1024
            print(f"         {pdf_path.name}  ({pdf_size} KB)")


if __name__ == "__main__":
    main()
